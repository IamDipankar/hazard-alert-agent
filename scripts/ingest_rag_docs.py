"""
Ingest RAG documents (DOCX) → chunk → embed → store in PostgreSQL with pgvector.

Usage:
    python -m scripts.ingest_rag_docs [--dir path/to/docx/files]
"""

import sys
import os
import uuid
import re

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from sqlalchemy import create_engine
from sqlalchemy.orm import Session
from backend.app.database import Base
from backend.app.models.rag import RAGDocument, RAGChunk

from docx import Document as DocxDocument


DEFAULT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Map filenames to doc types and default hazard types
DOC_CONFIG = {
    "Cyclone_RAG_bangla_rimes.docx": {"doc_type": "cyclone_rag", "default_hazard": "cyclone"},
    "Flood_RAG_Bangla.docx": {"doc_type": "flood_rag", "default_hazard": "flood"},
    "script_rimes.docx": {"doc_type": "script", "default_hazard": None},
}

# Hazard keywords for auto-detection
HAZARD_KEYWORDS = {
    "cyclone": ["ঘূর্ণিঝড়", "cyclone", "storm", "ঝড়"],
    "storm_surge": ["জলোচ্ছ্বাস", "storm surge", "surge"],
    "flood": ["বন্যা", "flood", "প্লাবন"],
    "flash_flood": ["আকস্মিক বন্যা", "flash flood"],
    "heavy_rainfall": ["বৃষ্টিপাত", "rainfall", "বৃষ্টি"],
    "waterlogging": ["জলাবদ্ধতা", "waterlogging"],
    "landslide": ["ভূমিধস", "landslide", "ধস"],
    "river_erosion": ["নদী ভাঙন", "erosion", "ভাঙন"],
}

# Phase keywords
PHASE_KEYWORDS = {
    "before": ["প্রস্তুতি", "আগে", "before", "preparation", "advance"],
    "during": ["চলাকালে", "during", "সময়е", "active"],
    "after": ["পরে", "after", "recovery", "উদ্ধার"],
}

# Action priority keywords
PRIORITY_KEYWORDS = {
    "life_safety": ["জীবন", "বাঁচান", "সরে যান", "আশ্রয়", "evacuation", "life", "safety"],
    "health": ["ওষুধ", "স্বাস্থ্য", "চিকিৎসা", "medicine", "health"],
    "asset": ["সম্পদ", "গবাদিপশু", "নৌকা", "জাল", "ফসল", "asset", "livestock"],
    "survey": ["প্রশ্ন", "জিজ্ঞাসা", "question", "survey", "ask"],
    "escalation": ["জরুরি", "urgent", "emergency", "বিপদ"],
}


def detect_metadata(text: str, heading: str, default_hazard: str | None) -> dict:
    """Auto-detect chunk metadata from content and heading."""
    text_lower = (text + " " + heading).lower()

    # Detect hazard type
    hazard = default_hazard
    for h_type, keywords in HAZARD_KEYWORDS.items():
        if any(k in text_lower for k in keywords):
            hazard = h_type
            break

    # Detect phase
    phase = None
    for p, keywords in PHASE_KEYWORDS.items():
        if any(k in text_lower for k in keywords):
            phase = p
            break

    # Detect action priority
    action_priority = None
    for ap, keywords in PRIORITY_KEYWORDS.items():
        if any(k in text_lower for k in keywords):
            action_priority = ap
            break

    # Detect housing type
    housing = None
    if any(k in text_lower for k in ["ঝুপড়ি", "jhupri"]):
        housing = "jhupri"
    elif any(k in text_lower for k in ["কাঁচা", "kacha"]):
        housing = "kacha"
    elif any(k in text_lower for k in ["আধাপাকা", "semi-pucca"]):
        housing = "semi-pucca"
    elif any(k in text_lower for k in ["পাকা", "pucca"]):
        housing = "pucca"

    # Detect livelihood
    livelihood = None
    if any(k in text_lower for k in ["জেলে", "fisherman", "মাছ ধরা"]):
        livelihood = "fisherman"
    elif any(k in text_lower for k in ["মাছ চাষ", "fish farm", "ঘের"]):
        livelihood = "fish_farmer"
    elif any(k in text_lower for k in ["কৃষক", "farmer", "চাষ"]):
        livelihood = "farmer"
    elif any(k in text_lower for k in ["লবণ", "salt"]):
        livelihood = "salt_farmer"

    # Detect vulnerability group
    vuln = None
    if any(k in text_lower for k in ["বয়স্ক", "elderly", "প্রবীণ"]):
        vuln = "elderly"
    elif any(k in text_lower for k in ["গর্ভবতী", "pregnant"]):
        vuln = "pregnancy"
    elif any(k in text_lower for k in ["প্রতিবন্ধী", "disability", "disabled"]):
        vuln = "disability"
    elif any(k in text_lower for k in ["শিশু", "children", "বাচ্চা"]):
        vuln = "children"

    return {
        "hazard_type": hazard,
        "phase": phase,
        "action_priority": action_priority,
        "housing_type": housing,
        "livelihood": livelihood,
        "vulnerability_group": vuln,
    }


def chunk_document(paragraphs: list[str], headings: list[str], target_words: int = 350, overlap_words: int = 50) -> list[dict]:
    """Chunk document into overlapping sections of ~target_words."""
    chunks = []
    current_chunk = []
    current_heading = ""
    current_word_count = 0

    for para, heading in zip(paragraphs, headings):
        words = para.split()
        if not words:
            continue

        if heading:
            current_heading = heading

        current_chunk.append(para)
        current_word_count += len(words)

        if current_word_count >= target_words:
            chunk_text = "\n".join(current_chunk)
            chunks.append({"text": chunk_text, "heading": current_heading, "word_count": current_word_count})

            # Keep overlap
            overlap_text = " ".join(chunk_text.split()[-overlap_words:])
            current_chunk = [overlap_text]
            current_word_count = overlap_words

    # Last chunk
    if current_chunk:
        chunk_text = "\n".join(current_chunk)
        chunks.append({"text": chunk_text, "heading": current_heading, "word_count": len(chunk_text.split())})

    return chunks


def get_embedding(text: str) -> list[float] | None:
    """Get OpenAI embedding for a text chunk."""
    try:
        from openai import OpenAI
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            return None

        client = OpenAI(api_key=api_key)
        response = client.embeddings.create(model="text-embedding-3-small", input=text[:8000])
        return response.data[0].embedding
    except Exception as e:
        print(f"  ⚠ Embedding failed: {e}")
        return None


def process_docx(filepath: str, doc_config: dict, session: Session):
    """Process a DOCX file into chunks and store in DB."""
    filename = os.path.basename(filepath)
    print(f"\nProcessing: {filename}")

    doc = DocxDocument(filepath)

    # Extract paragraphs with heading tracking
    paragraphs = []
    headings = []
    current_heading = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style and para.style.name and "Heading" in para.style.name:
            current_heading = text

        paragraphs.append(text)
        headings.append(current_heading)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text = " | ".join(cells)
                paragraphs.append(text)
                headings.append(current_heading)

    print(f"  Found {len(paragraphs)} paragraphs/rows")

    # Chunk
    chunks = chunk_document(paragraphs, headings)
    print(f"  Created {len(chunks)} chunks")

    # Create document record
    rag_doc = RAGDocument(
        id=uuid.uuid4(),
        filename=filename,
        doc_type=doc_config["doc_type"],
        total_chunks=len(chunks),
    )
    session.add(rag_doc)
    session.flush()

    # Create chunk records with embeddings
    for i, chunk in enumerate(chunks):
        metadata = detect_metadata(chunk["text"], chunk["heading"], doc_config.get("default_hazard"))
        embedding = get_embedding(chunk["text"])

        rag_chunk = RAGChunk(
            id=uuid.uuid4(),
            document_id=rag_doc.id,
            chunk_index=i,
            content=chunk["text"],
            heading=chunk["heading"],
            word_count=chunk["word_count"],
            hazard_type=metadata.get("hazard_type"),
            phase=metadata.get("phase"),
            action_priority=metadata.get("action_priority"),
            housing_type=metadata.get("housing_type"),
            livelihood=metadata.get("livelihood"),
            vulnerability_group=metadata.get("vulnerability_group"),
        )

        # Set embedding via raw SQL if available
        session.add(rag_chunk)
        print(f"  Chunk {i + 1}/{len(chunks)}: {chunk['word_count']} words, hazard={metadata.get('hazard_type')}, phase={metadata.get('phase')}")

    session.commit()
    print(f"  ✅ Stored {len(chunks)} chunks for {filename}")


def main(doc_dir: str = DEFAULT_DIR):
    from dotenv import load_dotenv
    load_dotenv()

    db_url = os.getenv("DATABASE_URL_SYNC", "postgresql://postgres:postgres@localhost:5432/hazard_alert")
    engine = create_engine(db_url)
    Base.metadata.create_all(engine)

    with Session(engine) as session:
        for filename, config in DOC_CONFIG.items():
            filepath = os.path.join(doc_dir, filename)
            if os.path.exists(filepath):
                process_docx(filepath, config, session)
            else:
                print(f"⚠ File not found: {filepath}")
                # Check in Guideline files subdirectory
                alt_path = os.path.join(doc_dir, "Guideline files", filename)
                if os.path.exists(alt_path):
                    process_docx(alt_path, config, session)

    print("\n✅ RAG ingestion complete")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--dir", default=DEFAULT_DIR, help="Directory containing DOCX files")
    args = parser.parse_args()
    main(args.dir)
